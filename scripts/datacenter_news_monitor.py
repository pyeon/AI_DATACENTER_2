"""
Datacenter News Monitor v10.4 - WITH TRANSLATION
✅ v10.3 기반 + 네이버 파파고 번역
"""

import yfinance as yf
import requests
import os
import json
from datetime import datetime, timedelta
import time
import warnings
from collections import defaultdict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import feedparser
from urllib.parse import quote
import re

warnings.filterwarnings('ignore')


# ============================================================================
# CONFIGURATION
# ============================================================================

TELEGRAM_BOT_TOKEN = os.environ.get('TELEGRAM_BOT_TOKEN')
TELEGRAM_CHAT_ID = os.environ.get('TELEGRAM_CHAT_ID')
NAVER_CLIENT_ID = os.environ.get('NAVER_CLIENT_ID', '')
NAVER_CLIENT_SECRET = os.environ.get('NAVER_CLIENT_SECRET', '')

STOCKS = [
    {'name': 'NVIDIA', 'ticker': 'NVDA', 'priority': 1, 'country': 'US', 
     'search_terms': ['NVIDIA AI', 'NVIDIA datacenter']},
    {'name': 'AMD', 'ticker': 'AMD', 'priority': 1, 'country': 'US', 
     'search_terms': ['AMD AI chip', 'AMD datacenter']},
    {'name': 'Intel', 'ticker': 'INTC', 'priority': 2, 'country': 'US', 
     'search_terms': ['Intel datacenter', 'Intel AI']},
    {'name': 'Super Micro', 'ticker': 'SMCI', 'priority': 1, 'country': 'US', 
     'search_terms': ['Super Micro AI server']},
    {'name': 'Broadcom', 'ticker': 'AVGO', 'priority': 1, 'country': 'US', 
     'search_terms': ['Broadcom AI chip']},
    {'name': 'Micron', 'ticker': 'MU', 'priority': 1, 'country': 'US', 
     'search_terms': ['Micron HBM', 'Micron memory']},
    {'name': 'SK Hynix', 'ticker': '000660.KS', 'priority': 1, 'country': 'KR', 
     'search_terms': ['SK하이닉스 HBM', 'SK하이닉스 AI']},
    {'name': 'Samsung', 'ticker': '005930.KS', 'priority': 1, 'country': 'KR', 
     'search_terms': ['삼성전자 반도체', '삼성전자 HBM']},
    {'name': 'LS ELECTRIC', 'ticker': '010120.KS', 'priority': 1, 'country': 'KR', 
     'search_terms': ['LS ELECTRIC 데이터센터']},
    {'name': 'Hanmi', 'ticker': '042700.KQ', 'priority': 2, 'country': 'KR', 
     'search_terms': ['한미반도체 AI']},
]

ENGLISH_KEYWORDS = {
    'high': ['AI', 'GPU', 'HBM', 'datacenter', 'data center', 'earnings', 'chip'],
    'medium': ['partnership', 'contract', 'launch', 'investment'],
}

KOREAN_KEYWORDS = {
    'high': ['AI', 'HBM', 'GPU', '데이터센터', '반도체', '실적', '수주'],
    'medium': ['파트너십', '계약', '투자', '출시'],
}


# ============================================================================
# TRANSLATION - NAVER PAPAGO
# ============================================================================

def translate_with_papago(text, max_length=4900):
    """
    네이버 파파고로 영문 → 한글 번역
    실패 시 원문 반환 (안전)
    """
    if not text or len(text.strip()) == 0:
        return text
    
    # 이미 한글이면 번역 안 함
    korean_chars = sum(1 for c in text if '가' <= c <= '힣')
    if len(text) > 0 and korean_chars / len(text) > 0.3:
        return text
    
    # Naver API 키 없으면 원문 반환
    if not NAVER_CLIENT_ID or not NAVER_CLIENT_SECRET:
        return text
    
    try:
        # 파파고는 5000자 제한
        text = text.strip()
        if len(text) > max_length:
            text = text[:max_length-3] + "..."
        
        url = "https://openapi.naver.com/v1/papago/n2mt"
        headers = {
            "X-Naver-Client-Id": NAVER_CLIENT_ID,
            "X-Naver-Client-Secret": NAVER_CLIENT_SECRET,
            "Content-Type": "application/x-www-form-urlencoded; charset=UTF-8"
        }
        data = {
            "source": "en",
            "target": "ko",
            "text": text
        }
        
        response = requests.post(url, headers=headers, data=data, timeout=10)
        
        if response.status_code == 200:
            result = response.json()
            translated = result.get('message', {}).get('result', {}).get('translatedText', '')
            if translated and len(translated.strip()) > 0:
                return translated
        
        # 번역 실패 시 원문 반환
        return text
        
    except Exception as e:
        # 에러 시에도 원문 반환 (안전)
        return text


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def load_seen_links():
    """Load previously seen news links"""
    try:
        with open('news_history.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
            return set(data.get('seen_links', []))
    except:
        return set()


def save_seen_links(links):
    """Save seen news links"""
    try:
        data = {
            'last_updated': datetime.now().isoformat(),
            'seen_links': list(links)
        }
        with open('news_history.json', 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    except:
        pass


def calculate_score(title, keywords_dict):
    """Calculate relevance score"""
    text = title.lower()
    score = 0
    matched_keywords = []
    
    for keyword in keywords_dict.get('high', []):
        if keyword.lower() in text:
            score += 10
            matched_keywords.append(keyword)
            break
    
    if score == 0:
        for keyword in keywords_dict.get('medium', []):
            if keyword.lower() in text:
                score += 6
                matched_keywords.append(keyword)
                break
    
    return max(score, 1), matched_keywords


# ============================================================================
# NEWS COLLECTION
# ============================================================================

def get_google_news_rss(search_term, seen_links):
    """Collect news using Google News RSS"""
    news_list = []
    
    try:
        encoded_term = quote(search_term)
        rss_url = f"https://news.google.com/rss/search?q={encoded_term}&hl=en-US&gl=US&ceid=US:en"
        
        feed = feedparser.parse(rss_url)
        
        if not feed.entries:
            return []
        
        week_ago = datetime.now() - timedelta(days=7)
        
        for entry in feed.entries[:20]:
            try:
                title = entry.get('title', '').strip()
                link = entry.get('link', '').strip()
                
                if not title or not link or len(title) < 10:
                    continue
                    
                if link in seen_links:
                    continue
                
                published = entry.get('published_parsed')
                if published:
                    pub_date = datetime(*published[:6])
                    if pub_date < week_ago:
                        continue
                else:
                    pub_date = datetime.now()
                
                summary = entry.get('summary', '').strip()
                summary = re.sub(r'<[^>]+>', '', summary)
                summary = re.sub(r'http[s]?://\S+', '', summary)
                summary = summary.strip()[:300]
                
                if len(summary) < 20:
                    summary = ''
                
                publisher = entry.get('source', {}).get('title', 'Google News')
                
                news_list.append({
                    'title': title,
                    'description': summary,
                    'link': link,
                    'publisher': publisher,
                    'date': pub_date,
                    'source': 'Google News'
                })
                seen_links.add(link)
                
            except:
                continue
        
    except Exception as e:
        print(f"      [ERROR] Google News: {str(e)}")
    
    return news_list


def get_us_news(search_terms, seen_links):
    """Get US company news - Google News only"""
    all_news = []
    
    for term in search_terms[:2]:
        news = get_google_news_rss(term, seen_links)
        all_news.extend(news)
        print(f"      [{term}] {len(news)} articles")
        time.sleep(1)
    
    return all_news


def get_naver_news(search_term, seen_links):
    """Get Korean news from Naver API"""
    news_list = []
    
    try:
        if not NAVER_CLIENT_ID or not NAVER_CLIENT_SECRET:
            return []
        
        url = "https://openapi.naver.com/v1/search/news.json"
        headers = {
            "X-Naver-Client-Id": NAVER_CLIENT_ID,
            "X-Naver-Client-Secret": NAVER_CLIENT_SECRET
        }
        params = {
            "query": search_term,
            "display": 20,
            "sort": "date"
        }
        
        response = requests.get(url, headers=headers, params=params, timeout=10)
        
        if response.status_code != 200:
            return []
        
        items = response.json().get('items', [])
        week_ago = datetime.now() - timedelta(days=7)
        
        for item in items:
            try:
                title = item.get('title', '').replace('<b>', '').replace('</b>', '').strip()
                link = item.get('originallink', item.get('link', '')).strip()
                
                if not title or not link or len(title) < 10:
                    continue
                    
                if link in seen_links:
                    continue
                
                try:
                    pub_date_str = item.get('pubDate', '')
                    pub_date = datetime.strptime(pub_date_str, '%a, %d %b %Y %H:%M:%S %z')
                    pub_date = pub_date.replace(tzinfo=None)
                    if pub_date < week_ago:
                        continue
                except:
                    pub_date = datetime.now()
                
                description = item.get('description', '').replace('<b>', '').replace('</b>', '').strip()
                publisher = 'Naver'
                if item.get('originallink'):
                    try:
                        publisher = item['originallink'].split('/')[2]
                    except:
                        pass
                
                news_list.append({
                    'title': title,
                    'description': description,
                    'link': link,
                    'publisher': publisher,
                    'date': pub_date,
                    'source': 'Naver API'
                })
                seen_links.add(link)
                
            except:
                continue
        
    except Exception as e:
        print(f"      [ERROR] Naver: {str(e)}")
    
    return news_list


# ============================================================================
# OUTPUT GENERATION
# ============================================================================

def create_docx_report(news_by_company, filename='outputs/news_report.docx'):
    """Create Word document report"""
    os.makedirs('outputs', exist_ok=True)
    
    doc = Document()
    
    title = doc.add_heading('Datacenter News Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    for company, news_list in news_by_company.items():
        flag = "KR" if news_list[0]['country'] == 'KR' else "US"
        doc.add_heading(f'[{flag}] {company}', level=1)
        
        for news in news_list:
            emoji = "HIGH" if news['score'] >= 10 else "MED" if news['score'] >= 6 else "LOW"
            
            para = doc.add_paragraph()
            para.add_run(f'[{emoji}] ').bold = True
            
            # 번역된 제목 또는 원본
            title_text = news.get('translated_title', news['title'])
            para.add_run(title_text).bold = True
            
            # 번역된 설명 또는 원본
            if news.get('translated_description'):
                doc.add_paragraph(f"Summary: {news['translated_description']}")
            elif news.get('description'):
                doc.add_paragraph(f"Summary: {news['description']}")
            
            hours = int((datetime.now() - news['date']).total_seconds() / 3600)
            time_str = "Now" if hours < 1 else f"{hours}h ago" if hours < 24 else f"{hours//24}d ago"
            
            doc.add_paragraph(f'Time: {time_str} | Source: {news["publisher"]} ({news["source"]})')
            
            link_para = doc.add_paragraph()
            link_para.add_run('Link: ').bold = True
            link_para.add_run(news['link'])
            
            doc.add_paragraph()
    
    doc.save(filename)
    return filename


def send_telegram_message(text):
    """Send text message to Telegram"""
    try:
        url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
        response = requests.post(url, data={"chat_id": TELEGRAM_CHAT_ID, "text": text}, timeout=10)
        return response.status_code == 200
    except:
        return False


def send_telegram_document(file_path, caption=''):
    """Send document file to Telegram"""
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': TELEGRAM_CHAT_ID, 'caption': caption}
            url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendDocument"
            response = requests.post(url, files=files, data=data, timeout=30)
            return response.status_code == 200
    except:
        return False


# ============================================================================
# MAIN
# ============================================================================

def main():
    """Main execution"""
    
    print("="*70)
    print("Datacenter News Monitor v10.4 - WITH TRANSLATION")
    print("  ✅ Stable base (v10.3)")
    print("  ✅ Naver Papago translation")
    print("  ✅ Fallback to original if translation fails")
    print("="*70)
    
    print("\n[CONFIG]")
    print(f"  Telegram: {'✓' if TELEGRAM_BOT_TOKEN and TELEGRAM_CHAT_ID else '✗'}")
    print(f"  Naver API: {'✓' if NAVER_CLIENT_ID and NAVER_CLIENT_SECRET else '✗'}")
    print(f"  Translation: {'✓ Enabled' if NAVER_CLIENT_ID and NAVER_CLIENT_SECRET else '✗ Disabled (no API key)'}")
    
    seen_links = load_seen_links()
    print(f"  Seen links: {len(seen_links)}")
    
    print("\n" + "="*70)
    print("PHASE 1: NEWS COLLECTION")
    print("="*70)
    
    all_news_by_company = defaultdict(list)
    stats = {'google': 0, 'naver': 0}
    
    for idx, stock in enumerate(STOCKS, 1):
        print(f"\n[{idx}/{len(STOCKS)}] {stock['name']} ({stock['country']})")
        
        if stock['country'] == 'US':
            news = get_us_news(stock.get('search_terms', []), seen_links)
            stats['google'] += len(news)
        else:
            news = []
            for term in stock.get('search_terms', [stock['name']]):
                naver_news = get_naver_news(term, seen_links)
                news.extend(naver_news)
                stats['naver'] += len(naver_news)
                print(f"      [{term}] {len(naver_news)} articles")
                time.sleep(0.3)
        
        keywords = KOREAN_KEYWORDS if stock['country'] == 'KR' else ENGLISH_KEYWORDS
        
        for news_item in news:
            score, matched = calculate_score(news_item['title'], keywords)
            news_item['score'] = score
            news_item['matched_keywords'] = matched
            news_item['company'] = stock['name']
            news_item['country'] = stock['country']
            all_news_by_company[stock['name']].append(news_item)
    
    save_seen_links(seen_links)
    
    print("\n" + "="*70)
    print("COLLECTION STATS")
    print("="*70)
    print(f"Google: {stats['google']}")
    print(f"Naver: {stats['naver']}")
    print(f"TOTAL: {sum(stats.values())}")
    
    # 상위 2개씩 선택
    filtered = {}
    for company, news_list in all_news_by_company.items():
        news_list.sort(key=lambda x: (x['score'], x['date']), reverse=True)
        filtered[company] = news_list[:2]
    
    final_count = sum(len(n) for n in filtered.values())
    print(f"Final (top 2 each): {final_count}")
    
    if final_count == 0:
        msg = f"📰 데이터센터 뉴스\n\n뉴스 없음\n\nGoogle: {stats['google']}\nNaver: {stats['naver']}"
        send_telegram_message(msg)
        return
    
    # ============================================================================
    # PHASE 2: TRANSLATION (NEW!)
    # ============================================================================
    
    print("\n" + "="*70)
    print("PHASE 2: TRANSLATION")
    print("="*70)
    
    if NAVER_CLIENT_ID and NAVER_CLIENT_SECRET:
        translation_count = 0
        
        for company, news_list in filtered.items():
            print(f"\n[{company}]")
            for idx, news in enumerate(news_list, 1):
                # 영문 기사만 번역 (US 기업)
                if news['country'] == 'US':
                    print(f"  [{idx}] Translating title...")
                    news['translated_title'] = translate_with_papago(news['title'], 300)
                    
                    if news.get('description'):
                        print(f"  [{idx}] Translating description...")
                        news['translated_description'] = translate_with_papago(news['description'], 200)
                    
                    translation_count += 1
                    time.sleep(0.5)  # 파파고 API 속도 제한 고려
                else:
                    # 한글 기사는 번역 안 함
                    news['translated_title'] = news['title']
                    news['translated_description'] = news.get('description', '')
        
        print(f"\nTranslated: {translation_count} articles")
    else:
        print("  Translation disabled (no Naver API key)")
        # 번역 없이 원문 사용
        for company, news_list in filtered.items():
            for news in news_list:
                news['translated_title'] = news['title']
                news['translated_description'] = news.get('description', '')
    
    print("\n" + "="*70)
    print("PHASE 3: MESSAGE GENERATION")
    print("="*70)
    
    messages = []
    msg = f"📰 데이터센터 뉴스\n{datetime.now().strftime('%Y-%m-%d %H:%M')}\n{'='*30}\n\n"
    msg += f"📊 기사: {final_count}개\n"
    msg += f"Google: {stats['google']} | Naver: {stats['naver']}\n\n"
    
    for company, news_list in sorted(filtered.items(), 
                                    key=lambda x: max(n['score'] for n in x[1]), 
                                    reverse=True):
        flag = "🇰🇷" if news_list[0]['country'] == 'KR' else "🇺🇸"
        section = f"{flag} {company}\n{'-'*30}\n\n"
        
        for news in news_list:
            emoji = "🔥" if news['score'] >= 10 else "📈"
            hours = int((datetime.now() - news['date']).total_seconds() / 3600)
            time_str = "방금" if hours < 1 else f"{hours}시간 전" if hours < 24 else f"{hours//24}일 전"
            
            # 번역된 제목 사용 (없으면 원본)
            title_text = news.get('translated_title', news['title'])
            section += f"{emoji} {title_text}\n\n"
            
            # 번역된 설명 사용 (있으면)
            if news.get('translated_description') and len(news['translated_description'].strip()) > 0:
                section += f"{news['translated_description']}\n\n"
            
            section += f"⏰ {time_str} | 📰 {news['publisher']}\n"
            section += f"🔗 {news['link']}\n\n"
        
        section += "="*30 + "\n\n"
        
        if len(msg + section) > 3500:
            messages.append(msg)
            msg = section
        else:
            msg += section
    
    if msg:
        messages.append(msg)
    
    print(f"Generated {len(messages)} messages")
    
    docx = create_docx_report(filtered, f"outputs/news_{datetime.now().strftime('%Y%m%d')}.docx")
    
    print("\n" + "="*70)
    print(f"PHASE 4: TELEGRAM DELIVERY")
    print("="*70)
    
    for idx, m in enumerate(messages, 1):
        send_telegram_message(m)
        print(f"  Message {idx}: Sent")
        time.sleep(1)
    
    send_telegram_document(docx, '📰 데이터센터 뉴스 리포트')
    print("  DOCX: Sent")
    
    print("\n" + "="*70)
    print("✅ COMPLETE")
    print("="*70)


if __name__ == "__main__":
    main()
